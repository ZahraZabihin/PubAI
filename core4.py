#core4.py
import os
import tempfile
from dotenv import load_dotenv
from langchain.document_loaders import DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore import InMemoryDocstore
from langchain.chat_models import ChatOpenAI
from langchain.prompts.chat import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
import fitz  # PyMuPDF
import faiss
import pickle
from docx import Document as DocxDocument
import re
import zipfile

load_dotenv()

class DocumentAnalyzer:
    def __init__(self):
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        if self.openai_api_key is None:
            raise ValueError("Did not find OPENAI_API_KEY, please add an environment variable `OPENAI_API_KEY` which contains it")
        
        self.prompt_template = ChatPromptTemplate.from_messages([
            SystemMessagePromptTemplate.from_template("You are a helpful assistant."),
            HumanMessagePromptTemplate.from_template(
                "Answer the question based only on the following context:\n\n{context}\n\n---\n\nQuestion: {question}"
            )
        ])
        
        self.INDEX_FILE_PATH = "faiss_index"
        self.METADATA_FILE_PATH = "faiss_metadata.pkl"

    def extract_text_from_pdf(self, uploaded_files, text_dir):
        if not os.path.exists(text_dir):
            os.makedirs(text_dir)

        for uploaded_file in uploaded_files:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                tmp_file.write(uploaded_file.read())
                pdf_path = tmp_file.name

            pdf_document = fitz.open(pdf_path)
            text_content = ""

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text_content += page.get_text("text")

            text_filename = os.path.join(text_dir, f'{os.path.splitext(uploaded_file.name)[0]}.txt')
            with open(text_filename, 'w') as text_file:
                text_file.write(text_content)

        return text_dir

    def generate_data_store(self, text_dir):
        documents = self.load_documents(text_dir)
        chunks = self.split_text(documents)
        self.save_to_faiss(chunks)

    def load_documents(self, text_dir):
        loader = DirectoryLoader(text_dir, glob="*.txt")
        documents = loader.load()
        return documents

    def split_text(self, documents):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            length_function=len,
            add_start_index=True,
        )
        chunks = text_splitter.split_documents(documents)
        return chunks

    def save_to_faiss(self, chunks):
        embeddings = OpenAIEmbeddings(openai_api_key=self.openai_api_key)
        vector_store = FAISS.from_documents(chunks, embeddings)

        faiss.write_index(vector_store.index, self.INDEX_FILE_PATH)

        with open(self.METADATA_FILE_PATH, 'wb') as f:
            pickle.dump({
                'index_to_docstore_id': vector_store.index_to_docstore_id,
                'docstore': vector_store.docstore._dict,
            }, f)

    def question_database(self, question_text):
        embedding_function = OpenAIEmbeddings(openai_api_key=self.openai_api_key)
        index = faiss.read_index(self.INDEX_FILE_PATH)

        with open(self.METADATA_FILE_PATH, 'rb') as f:
            metadata = pickle.load(f)

        docstore = InMemoryDocstore(metadata['docstore'])

        vector_store = FAISS(
            index=index,
            index_to_docstore_id=metadata['index_to_docstore_id'],
            docstore=docstore,
            embedding_function=embedding_function
        )

        results = vector_store.similarity_search_with_relevance_scores(question_text, k=25)
        if len(results) == 0 or results[0][1] < 0.75:
            return None

        context_text = "\n\n---\n\n".join([doc.page_content for doc, _score in results])
        
        messages = self.prompt_template.format_messages(
            context=context_text, question=question_text
        )
        
        model = ChatOpenAI(
            model_name="gpt-4o",
            openai_api_key=self.openai_api_key
        )
        
        response = model(messages)
        response_text = response.content

        sources = [doc.metadata.get("source", None) for doc, _score in results]
        formatted_response, unique_sources = self.format_response_with_references(response_text, sources)

        if not context_text.strip():
            formatted_response = "The context does not provide specific information for the question."

        return question_text, formatted_response, unique_sources

    def format_response_with_references(self, response_text, sources):
        sentences = response_text.split('. ')
        unique_sources = list(dict.fromkeys(sources))
        references = []
        formatted_sentences = []

        for i, sentence in enumerate(sentences):
            if i < len(unique_sources):
                reference = f" [{i+1}]"
                references.append(reference)
                formatted_sentences.append(sentence + reference)
            else:
                formatted_sentences.append(sentence)

        formatted_response = '. '.join(formatted_sentences) + '.'
        return formatted_response, unique_sources

    def save_responses_to_doc(self, responses):
        doc = DocxDocument()
        doc.add_heading('Question Responses', level=1)

        for i, (question, response, references) in enumerate(responses):
            doc.add_heading(f'Question {i+1}: {question}', level=2)
            doc.add_paragraph(response)
            doc.add_heading('References:', level=3)
            for j, reference in enumerate(references):
                if reference:
                    doc.add_paragraph(f"[{j+1}] {os.path.basename(reference)}")

        doc_path = "question_responses.docx"
        doc.save(doc_path)
        return doc_path

    def analyze_documents(self, uploaded_files, queries, disease_name):
        with tempfile.TemporaryDirectory() as temp_dir:
            self.extract_text_from_pdf(uploaded_files, temp_dir)
            self.generate_data_store(temp_dir)

            responses = []

            for question_key, question_template in queries.items():
                question_text = question_template.format(disease_name=disease_name)
                result = self.question_database(question_text)
                if result is not None:
                    responses.append(result)

            doc_path = self.save_responses_to_doc(responses)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp_file:
                with zipfile.ZipFile(tmp_file.name, 'w') as zipf:
                    zipf.write(doc_path, os.path.basename(doc_path))
                zip_file_path = tmp_file.name

        return zip_file_path
